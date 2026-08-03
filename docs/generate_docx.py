# -*- coding: utf-8 -*-
"""
generate_docx.py — 生成南京大学官方网页与学院网页清单 docx
输出到桌面

用法: python generate_docx.py
"""

import json
import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 项目根目录（docs 的上一级）
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.join(BASE_DIR, 'src')
DESKTOP_PATH = os.path.join(os.path.expanduser('~'), 'Desktop')
DOCX_PATH = os.path.join(DESKTOP_PATH, '{}.docx')


def load_sites():
    # 通过 node 读取 sites.js
    tmp = os.path.join(BASE_DIR, '.tmp_sites.json')
    node_cmd = f"require('{SRC_DIR.replace(chr(92), '/')}/sites.js');require('fs').writeFileSync('{tmp.replace(chr(92), '/')}',JSON.stringify(require('{SRC_DIR.replace(chr(92), '/')}/sites.js')))"
    subprocess.run(
        ['node', '-e', node_cmd],
        cwd=SRC_DIR, check=True
    )
    with open(tmp, 'r', encoding='utf-8') as f:
        data = json.load(f)
    os.remove(tmp)
    return data


def create_docx(title, sections, output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    total = sum(len(items) for _, items in sections)
    sub = doc.add_paragraph(f'南京大学  |  共 {total} 个页面  |  生成时间: 2026-07-22')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph('')

    for section_title, items in sections:
        if not items:
            continue
        sh = doc.add_heading(f'{section_title}（{len(items)} 个）', level=2)
        sh.runs[0].font.size = Pt(12)
        sh.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        hdr[0].text = '名称'
        hdr[1].text = '网址'
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell._tc.get_or_add_tcPr()
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1F497D')
            cell._tc.tcPr.append(shd)

        for item in items:
            row = table.add_row().cells
            row[0].text = item['name']
            row[1].text = item['url']
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph('')

    doc.save(output_path)
    print(f'已保存: {output_path}')


def main():
    data = load_sites()

    # 官方网页清单
    official_sections = [
        ('学校概况',     data.get('overview', {}).get('sites', [])),
        ('党群组织',     data.get('admin', {}).get('sites', [])),
        ('行政部门',     data.get('admin2', {}).get('sites', [])),
        ('公共服务单位', data.get('service', {}).get('sites', [])),
        ('专题网站',     data.get('special', {}).get('sites', [])),
        ('搜索门户',     data.get('search', {}).get('sites', [])),
    ]

    create_docx(
        '南京大学官方网页清单',
        official_sections,
        DOCX_PATH.format('南京大学官方网页清单')
    )

    # 学院清单：理工科 vs 文科
    all_academic = data.get('academic', {}).get('sites', []) + data.get('medical', {}).get('sites', [])

    scicollege = [s for s in all_academic if any(k in s['url'] for k in [
        'cs.', 'software.', 'ai.', 'ese.', 'eng.', 'hjxy.', 'es.', 'sgos.', 'as.',
        'nh.', 'life.', 'med.', 'sme.', 'math.', 'physics.', 'astronomy.',
        'is.', 'ise.', 'ic.', 'sser.', 'gcee.', 'ra.', 'futuretech.', 'frontier.', 'bme.', 'arch.', 'dii.',
        'chem.', 'sxsyj', 'essi', 'ias', 'imb', 'iddi', 'hnc', 'dafls', 'tyb', 'edu.'
    ])]
    artcollege = [s for s in all_academic if s not in scicollege]

    college_sections = [
        ('理工科院系', scicollege),
        ('文科院系', artcollege),
        ('医学院', data.get('medical', {}).get('sites', [])),
    ]

    create_docx(
        '南京大学院系网页清单',
        college_sections,
        DOCX_PATH.format('南京大学院系网页清单')
    )

    print('\n完成！')


if __name__ == '__main__':
    main()
