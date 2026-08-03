# -*- coding: utf-8 -*-
"""
生成自动信息检索模块完整教程 Word 文档
输出到项目根目录：<项目根目录>/自动信息检索模块_使用教程.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 输出到项目根目录（docs 的上一级）
DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '自动信息检索模块_使用教程.docx')

# 项目根目录提示（文档中显示给用户）
PROJECT_NAME = '自动信息检索模块'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=(0x1F, 0x49, 0x7D)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.bold = bold
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_code(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x27, 0x1A, 0x6E)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_color='1F497D'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1].cells
        bg = 'FFFFFF' if ri % 2 == 0 else 'EEF3FA'
        for ci, cell_text in enumerate(row_data):
            row[ci].text = str(cell_text)
            set_cell_bg(row[ci], bg)
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p


def add_note(doc, text):
    """添加灰色提示块"""
    p = doc.add_paragraph()
    run = p.add_run('📌 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def create_tutorial():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ========== 标题页 ==========
    doc.add_paragraph('')
    doc.add_paragraph('')
    title_h = doc.add_heading(PROJECT_NAME, level=1)
    title_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_h.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_heading('南京大学智能搜索爬虫 Agent', level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('完整使用教程 | 版本 1.0.0 | 2026-07-22')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ========== 目录 ==========
    add_heading(doc, '目录', level=1)
    toc_items = [
        '1. 项目简介',
        '2. 环境准备',
        '3. 快速开始',
        '4. 项目结构',
        '5. 核心模块详解',
        '6. Agent 使用详解',
        '7. 数据接口说明',
        '8. 命令行工具',
        '9. 数据存储',
        '10. 站点清单',
        '11. 常见问题',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.size = Pt(11)
    doc.add_page_break()

    # ========== 1. 项目简介 ==========
    add_heading(doc, '1. 项目简介', level=1)
    add_para(doc, f'{PROJECT_NAME}是一套完整的南京大学信息爬取与智能问答系统，包含 130 个站点的爬虫、智能搜索、AI Agent 问答、数据清洗与存储功能。', indent=True)

    add_heading(doc, '核心能力', level=2)
    add_table(doc, ['功能', '说明'], [
        ['站点爬取', '130 个官方站点的首页及内容'],
        ['智能搜索', '智搜门户 AI 搜索（需配置环境变量）'],
        ['自然语言问答', '输入问题，Agent 自动识别意图并回答'],
        ['数据清洗', '统一标题、内容、URL、时间、标签标准化'],
        ['数据导出', 'JSONL / JSON / CSV 多种格式'],
    ])
    doc.add_paragraph('')

    # ========== 2. 环境准备 ==========
    add_heading(doc, '2. 环境准备', level=1)
    add_heading(doc, '2.1 系统要求', level=2)
    add_table(doc, ['项目', '要求'], [
        ['操作系统', 'Windows 10/11、macOS 或 Linux'],
        ['Node.js', '版本 18 或以上（推荐 Node 20+）'],
        ['Python', '仅生成 docx 需要，版本 3.7+'],
        ['python-docx', 'pip install python-docx'],
    ])

    doc.add_paragraph('')
    add_heading(doc, '2.2 搜索功能的浏览器服务配置', level=2)
    add_para(doc, '搜索功能和完整版 Agent 需要浏览器服务。如果只需爬取静态页面，可跳过此步，使用 standalone/ 目录中的独立爬虫。', indent=True)

    add_heading(doc, '第一步：解压浏览器服务', level=3)
    add_para(doc, '找到本项目附带的压缩包"浏览器服务_skill.zip"，解压到合适位置，如 C:\\tools\\skills-nju-browser\\', indent=True)

    add_heading(doc, '第二步：配置环境变量', level=3)
    add_para(doc, 'Windows CMD（临时）：', indent=True)
    add_code(doc, 'set NJU_BROWSER_SKILLS=C:\\tools\\skills-nju-browser')
    add_para(doc, '永久配置：系统属性 → 高级 → 环境变量 → 新建系统变量 → 变量名：NJU_BROWSER_SKILLS', indent=True)

    add_heading(doc, '第三步：启动浏览器服务', level=3)
    add_code(doc, 'cd C:\\tools\\skills-nju-browser')
    add_code(doc, 'node nju-browser-start.js')
    add_note(doc, '如果跳过此步，仍可使用 standalone/ 目录中的独立爬虫，通过纯 HTTP 方式爬取静态页面。')

    doc.add_paragraph('')

    # ========== 3. 快速开始 ==========
    add_heading(doc, '3. 快速开始', level=1)

    add_heading(doc, '3.1 检查环境', level=2)
    add_para(doc, '双击运行项目根目录下的 配置检查.bat，确认 Node.js 已安装。', indent=True)

    add_heading(doc, '3.2 启动浏览器服务器（仅搜索需要）', level=2)
    add_para(doc, '如果只需要爬取站点功能，可以跳过此步。', indent=True)
    add_code(doc, 'cd <项目根目录>')
    add_code(doc, 'node %NJU_BROWSER_SKILLS%\\nju-browser-start.js')
    add_para(doc, '登录方式（二选一）：', indent=True)
    add_bullet(doc, '扫码登录（推荐）：用南京大学 App 或微信扫描二维码')
    add_bullet(doc, '账号密码登录：在浏览器窗口中输入学号和密码')

    add_heading(doc, '3.3 测试爬虫', level=2)
    add_para(doc, '打开命令行，进入项目根目录：', indent=True)
    add_code(doc, 'cd <项目根目录>')
    add_para(doc, '测试 1：爬取计算机学院首页', indent=True)
    add_code(doc, 'node src\\collector-generic.js --url https://cs.nju.edu.cn/')
    add_para(doc, '测试 2：启动 Agent 交互模式', indent=True)
    add_code(doc, 'node src\\agent\\index.js')

    add_heading(doc, '3.4 查看爬取结果', level=2)
    add_para(doc, '所有数据默认保存在 <项目根目录>/data/ 文件夹中。见本文档第 7 节「数据接口说明」。', indent=True)

    doc.add_page_break()

    # ========== 4. 项目结构 ==========
    add_heading(doc, '4. 项目结构', level=1)
    add_para(doc, '本项目包含两套系统：', indent=True)
    add_bullet(doc, 'standalone/ — 独立爬虫模块（浏览器登录 + 爬取 + 清洗 + 存储，不含 Agent）')
    add_bullet(doc, 'src/ — 完整系统（含 Agent 智能问答，需浏览器服务）')
    doc.add_paragraph('')
    add_table(doc, ['文件/目录', '说明'], [
        ['standalone/collector.js', '独立爬虫主程序（双模式：浏览器/HTTP）'],
        ['standalone/browser-start.js', '浏览器启动器（登录用）'],
        ['standalone/sites.js', '南京大学站点清单（130个）'],
        ['standalone/data/', '爬取结果目录（运行时创建）'],
        ['src/agent/', 'AI Agent（自然语言问答，含 Agent 才需要）'],
        ['src/crawler-core.js', '浏览器 API 引擎（Agent 系统使用）'],
        ['src/collector-generic.js', '页面采集器（需浏览器服务）'],
        ['src/cleaner.js', '数据清洗'],
        ['src/storage.js', '数据存储'],
        ['docs/generate_docx.py', '生成网页清单 Word 文档'],
        ['浏览器服务_skill.zip', '浏览器服务压缩包（需解压配置）'],
    ])
    doc.add_paragraph('')

    # ========== 5. 核心模块详解 ==========
    add_heading(doc, '5. 核心模块详解', level=1)

    add_heading(doc, '5.1 crawler-core.js — 核心引擎', level=2)
    add_table(doc, ['函数', '说明'], [
        ['api(method, path, data)', '底层 HTTP 调用'],
        ['navigate(url)', '导航到指定 URL'],
        ['extract()', '提取页面文本和链接'],
        ['evaluate(js)', '在页面中执行 JavaScript'],
        ['crawl(url, opts)', '爬取单个页面'],
        ['crawlBatch(urls, opts)', '批量爬取（支持并发控制）'],
        ['smartCrawl(url)', '智能爬取（自动选择适配器）'],
        ['findAdapter(url)', '根据 URL 找到对应适配器'],
    ])

    add_heading(doc, '5.2 站点适配器', level=2)
    add_table(doc, ['适配器', '网址', '说明'], [
        ['jw', 'jw.nju.edu.cn', '本科生院'],
        ['xsxy', 'xsxy.nju.edu.cn', '新生学院'],
        ['software', 'software.nju.edu.cn', '软件学院'],
        ['search-nju', 'search.nju.edu.cn', '智搜门户（支持搜索 API）'],
    ])

    add_heading(doc, '5.3 数据清洗', level=2)
    add_para(doc, '将原始数据统一清洗为标准格式，包含标题、内容、URL、发布时间、来源部门、标签等字段。')

    add_heading(doc, '5.4 数据存储', level=2)
    add_table(doc, ['格式', '说明', '方法'], [
        ['JSONL', '每行一条 JSON，方便追加', 'saveJSONL()'],
        ['JSON', '格式化 JSON 文件', 'saveJSON()'],
        ['CSV', '表格格式，可用 Excel 打开', 'saveCSV()'],
    ])
    add_para(doc, '数据目录：<项目根目录>/data/', indent=True)
    doc.add_paragraph('')

    # ========== 6. Agent 使用详解 ==========
    add_heading(doc, '6. Agent 使用详解', level=1)

    add_heading(doc, '6.1 Agent 交互模式（推荐新手）', level=2)
    add_code(doc, 'cd <项目根目录>')
    add_code(doc, 'node src\\agent\\index.js')
    add_para(doc, '输入自然语言问题，Agent 自动判断意图并爬取相关站点。', indent=True)
    add_para(doc, '示例对话：', indent=True)
    add_code(doc, '自动信息检索模块> 计算机学院')
    add_code(doc, '# → 返回计算机学院网页摘要')
    add_code(doc, '自动信息检索模块> 医学院')
    add_code(doc, '# → 返回医学院及6所附属医院信息')
    add_code(doc, '自动信息检索模块> 列出所有站点')
    add_code(doc, '# → 显示全部 130 个站点的分类目录')
    add_code(doc, '自动信息检索模块> 帮助')
    add_code(doc, '# → 显示功能说明')

    add_heading(doc, '6.2 LLM 模式配置', level=2)
    add_para(doc, 'Agent 支持三种运行模式，通过环境变量 LLM_MODE 切换：', indent=True)
    add_table(doc, ['模式', '说明', '是否需要 API Key'], [
        ['local', '本地模式：关键词匹配+模板，无需网络', '否'],
        ['claude', 'Claude API 模式：接入 Claude 大语言模型', '是（需 ANTHROPIC_API_KEY）'],
        ['openai', 'OpenAI 兼容模式：接入 GPT 等模型', '是（需 OPENAI_API_KEY）'],
    ])

    add_heading(doc, '配置 Claude API', level=3)
    add_para(doc, '1. 前往 Anthropic Console (console.anthropic.com) 注册并获取 API Key', indent=True)
    add_para(doc, '2. 配置环境变量（Windows）：', indent=True)
    add_code(doc, 'set ANTHROPIC_API_KEY=sk-ant-xxxxx-your-key-here')
    add_code(doc, 'set LLM_MODE=claude')
    add_para(doc, '3. 重启 Agent：', indent=True)
    add_code(doc, 'node src\\agent\\index.js')

    add_heading(doc, '配置 OpenAI API', level=3)
    add_code(doc, 'set OPENAI_API_KEY=sk-xxxxx-your-key-here')
    add_code(doc, 'set LLM_MODE=openai')

    add_heading(doc, '永久保存配置', level=3)
    add_para(doc, 'Windows：系统 → 高级 → 环境变量 → 在系统变量中新建', indent=True)
    add_para(doc, 'Linux/macOS：在 ~/.bashrc 或 ~/.zshrc 中添加 export 语句', indent=True)

    add_heading(doc, '6.3 LLM 模式切换效果对比', level=2)
    add_table(doc, ['输入', 'local 模式响应', 'claude 模式响应'], [
        ['"计算机学院"', '返回站点摘要', '返回智能分析后的总结'],
        ['"南京大学有什么优势学科？"', '引导用户使用搜索功能', '综合多个站点给出完整回答'],
        ['"搜索 机器学习"', '调用搜索 API 返回结果', '同左（搜索功能不依赖 LLM）'],
    ])

    add_heading(doc, '6.4 Agent 内部工作流程', level=2)
    add_para(doc, '用户输入 → 意图识别 → 关键词站点匹配/slots 搜索API → 爬取页面 → 格式化输出 → 返回给用户', indent=True)

    doc.add_page_break()

    # ========== 7. 数据接口说明 ==========
    add_heading(doc, '7. 数据接口说明', level=1)

    add_heading(doc, '7.1 数据输出位置', level=2)
    add_para(doc, '所有爬取结果保存在 <项目根目录>/data/ 目录：', indent=True)
    add_code(doc, 'data/')
    add_code(doc, '├── all_records.jsonl       # 全量记录（追加去重，主文件）')
    add_code(doc, '├── crawl_2026-07-22_14-30-00.jsonl   # 分次记录（带时间戳）')
    add_code(doc, '└── crawl_2026-07-22_15-00-00.csv     # CSV 格式')

    add_heading(doc, '7.2 数据格式示例', level=2)
    add_para(doc, 'JSONL 格式（每行一条，可用任何文本编辑器查看）：', indent=True)
    add_code(doc, '{"title":"计算机学院主页","url":"https://cs.nju.edu.cn/","content":"...","publishTime":"2026-07-22","source":{"department":"计算机学院"},"tags":["通知公告"],"crawler":"nju-crawler","crawlTime":"2026-07-22T14:30:00.000Z"}')

    add_heading(doc, '7.3 在代码中使用数据', level=2)
    add_para(doc, '读取所有记录：', indent=True)
    add_code(doc, "const { readAllJSONL } = require('./src/storage');")
    add_code(doc, 'const records = readAllJSONL();')
    add_code(doc, 'console.log(`共 ${records.length} 条记录`);')
    add_para(doc, '统计数据量：', indent=True)
    add_code(doc, 'node -e "console.log(require(\'./src/storage\').countRecords())"')

    add_heading(doc, '7.4 数据字段说明', level=2)
    add_table(doc, ['字段', '类型', '说明'], [
        ['title', 'string', '页面标题'],
        ['url', 'string', '页面 URL'],
        ['content', 'string', '页面正文（去 HTML 标签）'],
        ['publishTime', 'string', '标准化日期（YYYY-MM-DD）'],
        ['source.department', 'string', '推断的部门名称'],
        ['source.siteName', 'string', '站点显示名称'],
        ['tags', 'string[]', '自动推断的标签数组'],
        ['crawler', 'string', '固定值 "nju-crawler"'],
        ['crawlTime', 'string', 'ISO 格式爬取时间'],
    ])
    doc.add_paragraph('')

    # ========== 8. 命令行工具 ==========
    add_heading(doc, '8. 命令行工具', level=1)

    add_heading(doc, '8.1 collector-search.js — 搜索采集器', level=2)
    add_code(doc, 'node src\\collector-search.js --keyword 机器学习 --max 20 --output')
    add_table(doc, ['参数', '说明', '示例'], [
        ['--keyword', '搜索关键词', '--keyword 机器学习'],
        ['--type', '搜索类型', 'zh/xwzx/xsbg/jszy'],
        ['--max', '最大结果数', '--max 50'],
        ['--output', '保存到文件', '--output'],
        ['--format', '输出格式', '--format csv'],
    ])

    add_heading(doc, '8.2 collector-generic.js — 通用采集器', level=2)
    add_table(doc, ['参数', '说明', '示例'], [
        ['--url URL', '爬取单个页面', '--url https://cs.nju.edu.cn/'],
        ['--all', '爬取全部站点', '--all'],
        ['--list FILE', '从文件读取 URL 列表', '--list urls.txt'],
        ['--output', '保存结果', '--output'],
        ['--deduplicate', '去重追加', '--deduplicate'],
    ])

    add_heading(doc, '8.3 run-crawl.js — 一键全量采集', level=2)
    add_code(doc, 'node src\\run-crawl.js')

    add_heading(doc, '8.4 generate_docx.py — 生成网页清单', level=2)
    add_code(doc, 'pip install python-docx')
    add_code(doc, 'cd <项目根目录>')
    add_code(doc, 'python docs\\generate_docx.py')

    doc.add_page_break()

    # ========== 9. 数据存储 ==========
    add_heading(doc, '9. 数据存储', level=1)
    add_para(doc, '采集数据保存在：<项目根目录>/data/', indent=True)
    add_heading(doc, '查看数据', level=2)
    add_code(doc, 'node -e "console.log(require(\'./src/storage\').countRecords())"')

    # ========== 10. 站点清单 ==========
    add_heading(doc, '10. 站点清单（共 130 个页面）', level=1)
    add_table(doc, ['分类', '数量', '说明'], [
        ['学校概况', '5', '南大简介、现任领导、历任领导、校史、标识'],
        ['党群组织', '16', '党委办、纪委、组织部、宣传部、统战部等'],
        ['行政部门', '26', '校长办、人力处、科研院、研究生院等'],
        ['教学科研单位', '56', '全部学院（含苏州校区）、研究所、研究中心'],
        ['医学院', '7', '医学院+6所附属医院'],
        ['公共服务单位', '15', '图书馆、档案馆、医院、超算中心等'],
        ['专题网站', '4', '信息公开、人才招聘、招生'],
        ['搜索门户', '1', '南京大学智搜门户'],
    ])
    doc.add_paragraph('')

    # ========== 11. 常见问题 ==========
    add_heading(doc, '11. 常见问题', level=1)
    qa_data = [
        ('Q1：搜索返回 0 条？', '浏览器会话过期。关闭后重新启动服务器并登录。\ncurl -X POST http://127.0.0.1:4100/shutdown\n然后重新运行 nju-browser-start.js'),
        ('Q2：爬取页面报错？', '可能是网络临时故障，稍后重试。'),
        ('Q3：端口 4100 被占用？', '先关闭旧服务器：curl -X POST http://127.0.0.1:4100/shutdown'),
        ('Q4：如何查看采集数量？', 'node -e "console.log(require(\'./src/storage\').countRecords())"'),
        ('Q5：如何添加新站点？', '在 src/site-adapters/ 新建适配器文件，然后在 index.js 注册。'),
        ('Q6：Agent 如何接入 LLM？', '设置环境变量后重启：\nset ANTHROPIC_API_KEY=your_key\nset LLM_MODE=claude'),
        ('Q7：python-docx 未安装？', 'pip install python-docx'),
        ('Q8：找不到 skills-nju-browser？', '这说明您未安装 Claude Code 或该 skill。\n这不是错误——站点爬取功能完全不需要此服务！'),
    ]
    for q, a in qa_data:
        add_heading(doc, q, level=3)
        p = doc.add_paragraph(a)
        for run in p.runs:
            run.font.size = Pt(10.5)

    # ========== 附录 ==========
    add_heading(doc, '附录：快速命令索引', level=1)
    add_table(doc, ['命令', '功能'], [
        ['node src\\agent\\index.js', '启动 Agent 交互'],
        ['node src\\collector-search.js --keyword XXX --max 20 --output', '搜索采集'],
        ['node src\\collector-generic.js --url URL', '爬取单个页面'],
        ['node src\\collector-generic.js --all', '爬取全部 130 个站点'],
        ['node src\\run-crawl.js', '一键全量采集'],
        ['python docs\\generate_docx.py', '生成网页清单 docx'],
        ['node -e "console.log(require(\'./src/storage\').countRecords())"', '查看数据量'],
    ])

    doc.save(DOCX_PATH)
    print(f'已保存: {DOCX_PATH}')


if __name__ == '__main__':
    create_tutorial()
